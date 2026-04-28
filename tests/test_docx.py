import sys, os
import tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))

from processors.docx_processor import DOCXProcessor
from docx import Document


def make_test_docx(path):
    doc = Document()
    doc.add_heading("Test Heading", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Another sentence here.")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "Value"
    t.rows[1].cells[0].text = "Ram"
    t.rows[1].cells[1].text = "Engineer"
    doc.save(path)


def test_extract_docx():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    make_test_docx(path)

    proc = DOCXProcessor()
    result = proc.extract(path)

    texts = [u["text"] for u in result["translation_units"]]
    assert "This is a test paragraph." in texts
    assert "Ram" in texts
    assert "Name" in texts

    os.unlink(path)


def test_rebuild_docx():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        in_path = f.name
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        out_path = f.name

    make_test_docx(in_path)

    proc = DOCXProcessor()
    data = proc.extract(in_path)

    for unit in data["translation_units"]:
        unit["translated"] = "[TRANSLATED] " + unit["text"]
        unit["sentences"] = [unit["text"]]

    stats = proc.rebuild(data, out_path)

    rebuilt = Document(out_path)
    all_text = " ".join(p.text for p in rebuilt.paragraphs)
    assert "[TRANSLATED]" in all_text

    os.unlink(in_path)
    os.unlink(out_path)