from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import copy


class DOCXProcessor:

    def extract(self, docx_path: str) -> dict:
        doc = Document(docx_path)
        units = []

        # 1. Body paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                units.append({
                    "text": text,
                    "location": {"type": "body", "para_idx": i},
                    "translated": "",
                    "sentences": [],
                })

        # 2. Tables
        for t_i, table in enumerate(doc.tables):
            for r_i, row in enumerate(table.rows):
                for c_i, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        units.append({
                            "text": text,
                            "location": {
                                "type": "table",
                                "table_idx": t_i,
                                "row_idx": r_i,
                                "col_idx": c_i,
                            },
                            "translated": "",
                            "sentences": [],
                        })

        # 3. Headers & Footers
        for sec_i, section in enumerate(doc.sections):
            for region, rtype in [
                (section.header, "header"),
                (section.footer, "footer"),
            ]:
                if region:
                    for p_i, para in enumerate(region.paragraphs):
                        text = para.text.strip()
                        if text:
                            units.append({
                                "text": text,
                                "location": {
                                    "type": rtype,
                                    "section_idx": sec_i,
                                    "para_idx": p_i,
                                },
                                "translated": "",
                                "sentences": [],
                            })

        # 4. Footnotes (if present)
        try:
            fn_part = doc.part.footnotes_part
            if fn_part:
                for fn_id, fn in fn_part.footnotes.items():
                    for p_i, para in enumerate(fn.paragraphs):
                        text = para.text.strip()
                        if text:
                            units.append({
                                "text": text,
                                "location": {
                                    "type": "footnote",
                                    "fn_id": fn_id,
                                    "para_idx": p_i,
                                },
                                "translated": "",
                                "sentences": [],
                            })
        except Exception:
            pass

        return {
            "source_path": docx_path,
            "translation_units": units,
        }

    def rebuild(self, doc_data: dict, output_path: str, direction: str = "en→ne") -> dict:
        source_path = doc_data["source_path"]
        units = doc_data["translation_units"]
        doc = Document(source_path)

        translated_count = 0

        for unit in units:
            translated = unit.get("translated", "").strip()
            if not translated or translated == unit["text"]:
                continue

            loc = unit["location"]
            ltype = loc["type"]

            try:
                if ltype == "body":
                    para = doc.paragraphs[loc["para_idx"]]
                    self._apply_to_paragraph(para, unit["text"], translated)
                    translated_count += 1

                elif ltype == "table":
                    table = doc.tables[loc["table_idx"]]
                    cell = table.rows[loc["row_idx"]].cells[loc["col_idx"]]
                    for para in cell.paragraphs:
                        if para.text.strip() == unit["text"]:
                            self._apply_to_paragraph(para, unit["text"], translated)
                            translated_count += 1
                            break

                elif ltype in ("header", "footer"):
                    section = doc.sections[loc["section_idx"]]
                    region = section.header if ltype == "header" else section.footer
                    if region:
                        para = region.paragraphs[loc["para_idx"]]
                        self._apply_to_paragraph(para, unit["text"], translated)
                        translated_count += 1

            except (IndexError, KeyError, AttributeError):
                pass

        doc.save(output_path)

        failed = len(units) - translated_count
        return {
            "translated": translated_count,
            "failed": failed,
            "total_sentences": len(units),
        }

    def _apply_to_paragraph(self, para, original: str, translated: str):
        """
        Replace paragraph text while preserving run-level formatting.
        Strategy: distribute translated text proportionally across runs.
        """
        runs = para.runs
        if not runs:
            return

        if len(runs) == 1:
            runs[0].text = translated
            return

        # Proportional distribution across runs
        total_orig_len = sum(len(r.text) for r in runs)
        trans_len = len(translated)
        char_pos = 0

        for i, run in enumerate(runs):
            if i == len(runs) - 1:
                run.text = translated[char_pos:]
            else:
                ratio = len(run.text) / max(total_orig_len, 1)
                chars = max(1, int(ratio * trans_len))
                end = min(char_pos + chars, len(translated))
                # Cut at word boundary
                while end < len(translated) and translated[end] != " ":
                    end += 1
                run.text = translated[char_pos:end]
                char_pos = end